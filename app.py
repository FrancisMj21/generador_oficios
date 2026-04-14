from datetime import datetime
from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from dotenv import load_dotenv
from flask import Flask, jsonify, redirect, render_template, request, send_file
from supabase import create_client

import os
import re


load_dotenv()

SUPABASE_URL = os.getenv("SUPABASE_URL")
SUPABASE_KEY = os.getenv("SUPABASE_KEY")

supabase = create_client(SUPABASE_URL, SUPABASE_KEY)

app = Flask(__name__)

BASE_DIR = Path(__file__).resolve().parent
TEMPLATES_DOCX_DIR = BASE_DIR / "plantillas"
GENERATED_DIR = BASE_DIR / "generados"
GENERATED_DIR.mkdir(exist_ok=True)


def normalizar_texto(valor):
    return str(valor or "").strip()


def slugify(valor):
    texto = normalizar_texto(valor).lower()
    texto = re.sub(r"[^a-z0-9]+", "_", texto)
    return texto.strip("_") or "archivo"


def obtener_lista(nombre_tabla):
    response = supabase.table(nombre_tabla).select("*").execute()
    return response.data or []


def enriquecer_personas(personas, tipos_persona, condiciones):
    tipo_map = {tipo["id"]: tipo["nombre"] for tipo in tipos_persona}
    condicion_map = {cond["id"]: cond["nombre"] for cond in condiciones}

    for persona in personas:
        persona["tipo_persona"] = tipo_map.get(
            persona.get("tipo_persona_id"),
            persona.get("tipo_persona", ""),
        )
        persona["condicion"] = condicion_map.get(
            persona.get("condicion_id"),
            persona.get("condicion", ""),
        )
        persona["estado"] = normalizar_texto(persona.get("estado")) or "Pendiente"

    return personas


def obtener_personas():
    personas = obtener_lista("personas")
    tipos_persona = obtener_lista("tipos_persona")
    condiciones = obtener_lista("condiciones")
    return enriquecer_personas(personas, tipos_persona, condiciones), tipos_persona, condiciones


def obtener_persona_por_id(persona_id):
    response = supabase.table("personas").select("*").eq("id", persona_id).limit(1).execute()
    data = response.data or []
    if not data:
        return None

    persona = data[0]
    tipos_persona = obtener_lista("tipos_persona")
    condiciones = obtener_lista("condiciones")
    return enriquecer_personas([persona], tipos_persona, condiciones)[0]


def nombre_archivo_oficio(persona):
    return f"oficio_{persona['id']}_{slugify(persona.get('nombre'))}.docx"


def ruta_archivo_generado(persona):
    return GENERATED_DIR / nombre_archivo_oficio(persona)


def obtener_lineas_periodos(persona):
    periodos = normalizar_texto(persona.get("periodos"))
    if not periodos:
        return []
    return [linea.strip() for linea in periodos.splitlines() if linea.strip()]


def obtener_fecha_actual():
    meses = {
        1: "enero",
        2: "febrero",
        3: "marzo",
        4: "abril",
        5: "mayo",
        6: "junio",
        7: "julio",
        8: "agosto",
        9: "septiembre",
        10: "octubre",
        11: "noviembre",
        12: "diciembre",
    }
    ahora = datetime.now()
    return f"{ahora.day} de {meses[ahora.month]} de {ahora.year}"


def numero_oficio(persona):
    return f"{persona['id']}-{datetime.now().year}"


def tipo_template(persona):
    return "docente.docx"


def plantilla_disponible(persona):
    candidate = TEMPLATES_DOCX_DIR / tipo_template(persona)
    if candidate.exists() and candidate.stat().st_size > 0:
        return candidate

    fallback = TEMPLATES_DOCX_DIR / "docente.docx"
    if fallback.exists() and fallback.stat().st_size > 0:
        return fallback

    return None


def reemplazar_en_parrafos(parrafos, reemplazos):
    for parrafo in parrafos:
        texto = parrafo.text
        nuevo = texto
        for origen, destino in reemplazos.items():
            if origen in nuevo:
                nuevo = nuevo.replace(origen, destino)
        if nuevo != texto:
            for run in parrafo.runs:
                run.text = ""
            if parrafo.runs:
                parrafo.runs[0].text = nuevo
            else:
                parrafo.add_run(nuevo)


def reemplazar_en_tablas(tablas, reemplazos):
    for tabla in tablas:
        for fila in tabla.rows:
            for celda in fila.cells:
                reemplazar_en_parrafos(celda.paragraphs, reemplazos)
                reemplazar_en_tablas(celda.tables, reemplazos)


def aplicar_reemplazos(documento, reemplazos):
    reemplazar_en_parrafos(documento.paragraphs, reemplazos)
    reemplazar_en_tablas(documento.tables, reemplazos)

    for section in documento.sections:
        reemplazar_en_parrafos(section.header.paragraphs, reemplazos)
        reemplazar_en_tablas(section.header.tables, reemplazos)
        reemplazar_en_parrafos(section.footer.paragraphs, reemplazos)
        reemplazar_en_tablas(section.footer.tables, reemplazos)


def construir_reemplazos(persona):
    nombre = normalizar_texto(persona.get("nombre")).upper()
    condicion = normalizar_texto(persona.get("condicion")).upper()
    cargo_actual = normalizar_texto(persona.get("cargo_actual"))
    if not cargo_actual:
        tipo = normalizar_texto(persona.get("tipo_persona")).upper()
        cargo_actual = f"{tipo} {condicion}".strip()

    centro = normalizar_texto(persona.get("centro_trabajo")).upper()
    direccion = normalizar_texto(persona.get("direccion"))
    provincia = normalizar_texto(persona.get("provincia")) or "TACNA"
    telefono = normalizar_texto(persona.get("telefono"))
    correo = normalizar_texto(persona.get("correo"))
    solicita = normalizar_texto(persona.get("solicita"))
    fecha = obtener_fecha_actual()
    oficio = numero_oficio(persona)

    return {
        "{{FECHA}}": fecha,
        "{{OFICIO}}": oficio,
        "{{NOMBRE}}": nombre,
        "{{CARGO}}": cargo_actual.upper(),
        "{{CENTRO_TRABAJO}}": centro,
        "{{DIRECCION}}": direccion,
        "{{PROVINCIA}}": provincia,
        "{{TELEFONO}}": telefono,
        "{{CORREO}}": correo,
        "{{SOLICITA}}": solicita,
        "Tacna, ": f"Tacna, {fecha}",
        "KARINA PAREDES NINAJA": nombre,
        "PROFESOR NOMBRADO": cargo_actual.upper(),
        "I.E. JUAN VELASCO ALVARADO": centro,
        "DIRECCION: CPM Leguia Mz 28 Lte 15.": f"DIRECCION: {direccion}.",
        "Tacna/Tacna/Tacna": f"{provincia}/{provincia}/{provincia}",
        "995968988": telefono,
        "karinaparedes529@gmail.com": correo,
        "reconocimiento de años de servicio": solicita,
        "OFICIO N°      - 2026-UFESC-URRHH/UGEL.T/GOB.REG.TACNA": f"OFICIO N° {oficio}-UFESC-URRHH/UGEL.T/GOB.REG.TACNA",
        "OFICIO N°": f"OFICIO N° {oficio}",
    }


def documento_tiene_placeholders(documento):
    contenido = "\n".join(parrafo.text for parrafo in documento.paragraphs)
    return "{{" in contenido and "}}" in contenido


def construir_documento_desde_cero(persona):
    documento = Document()

    fecha = obtener_fecha_actual()
    oficio = numero_oficio(persona)
    nombre = normalizar_texto(persona.get("nombre")).upper()
    condicion = normalizar_texto(persona.get("condicion")).upper()
    tipo = normalizar_texto(persona.get("tipo_persona")).upper()
    cargo_actual = normalizar_texto(persona.get("cargo_actual")) or f"{tipo} {condicion}"
    centro = normalizar_texto(persona.get("centro_trabajo")).upper()
    direccion = normalizar_texto(persona.get("direccion"))
    provincia = normalizar_texto(persona.get("provincia")) or "TACNA"
    telefono = normalizar_texto(persona.get("telefono"))
    correo = normalizar_texto(persona.get("correo"))
    solicita = normalizar_texto(persona.get("solicita"))
    periodos = obtener_lineas_periodos(persona)

    documento.add_paragraph(f"Tacna, {fecha}")
    documento.add_paragraph("")
    documento.add_paragraph(f"OFICIO N° {oficio}-UFESC-URRHH/UGEL.T/GOB.REG.TACNA")
    documento.add_paragraph("")
    documento.add_paragraph("SEÑOR:")
    documento.add_paragraph(nombre)
    documento.add_paragraph(cargo_actual.upper())
    documento.add_paragraph(centro)
    documento.add_paragraph(f"DIRECCIÓN: {direccion}")
    documento.add_paragraph(f"{provincia}/{provincia}/{provincia}")
    documento.add_paragraph(f"Teléfono: {telefono}")
    documento.add_paragraph(f"Correo electrónico: {correo}")
    documento.add_paragraph("")
    documento.add_paragraph("ASUNTO: FORMULO RESPUESTA")
    documento.add_paragraph("REFERENCIA: CUD. N° __________")
    documento.add_paragraph("")

    documento.add_paragraph(
        "Tengo el agrado de dirigirme a usted, para expresarle mi cordial saludo y a la vez "
        f"brindar la debida atención al documento de la referencia, mediante el cual solicita {solicita}."
    )
    documento.add_paragraph(
        "Al respecto, se procedió a evaluar el expediente de la referencia, en el que se pudo "
        "verificar que falta presentar la documentación sustentatoria de los periodos indicados. "
        "Por lo que se exhorta cumplir con la subsanación correspondiente para continuar con el "
        "procedimiento administrativo, bajo su responsabilidad."
    )

    if periodos:
        documento.add_paragraph("Periodos observados:")
        for periodo in periodos:
            documento.add_paragraph(periodo, style="List Bullet")

    documento.add_paragraph("")
    documento.add_paragraph(
        "Sin otro particular, hago propicia la oportunidad para reiterarle las muestras de mi "
        "especial consideración y estima personal."
    )
    documento.add_paragraph("")
    documento.add_paragraph("Atentamente,")
    documento.add_paragraph("")
    documento.add_paragraph("GOBIERNO REGIONAL DE TACNA")
    documento.add_paragraph("____________________________________________")
    documento.add_paragraph("JEFE DE LA UNIDAD DE RECURSOS HUMANOS")
    documento.add_paragraph("UGEL TACNA")

    return documento


def poblar_documento(documento, persona):
    reemplazos = construir_reemplazos(persona)
    aplicar_reemplazos(documento, reemplazos)

    periodos = obtener_lineas_periodos(persona)
    if periodos:
        documento.add_paragraph("")
        documento.add_paragraph("Periodos observados:")
        for periodo in periodos:
            documento.add_paragraph(periodo, style="List Bullet")

    return documento


def generar_documento_persona(persona):
    plantilla = plantilla_disponible(persona)

    if plantilla:
        documento = Document(str(plantilla))
        if documento_tiene_placeholders(documento):
            aplicar_reemplazos(documento, construir_reemplazos(persona))
        else:
            documento = poblar_documento(documento, persona)
    else:
        documento = construir_documento_desde_cero(persona)

    salida = ruta_archivo_generado(persona)
    documento.save(str(salida))
    return salida


def actualizar_estado_generado(persona_id):
    try:
        supabase.table("personas").update({"estado": "Generado"}).eq("id", persona_id).execute()
    except Exception:
        pass


def generar_y_marcar(persona):
    archivo = generar_documento_persona(persona)
    actualizar_estado_generado(persona["id"])
    return archivo


@app.route("/")
def index():
    personas, tipos_persona, condiciones = obtener_personas()

    total = len(personas)
    generados = sum(1 for p in personas if str(p.get("estado", "")).lower() == "generado")
    pendientes = total - generados

    return render_template(
        "index.html",
        personas=personas,
        tipos_persona=tipos_persona,
        condiciones=condiciones,
        total=total,
        generados=generados,
        pendientes=pendientes,
    )


@app.route("/guardar", methods=["POST"])
def guardar():
    data = {
        "dni": request.form.get("dni", ""),
        "nombre": request.form["nombre"],
        "tipo_persona_id": request.form["tipo_persona_id"],
        "cargo_actual": request.form.get("cargo_actual", ""),
        "condicion_id": request.form["condicion_id"],
        "centro_trabajo": request.form["centro_trabajo"],
        "direccion": request.form.get("direccion", ""),
        "provincia": request.form.get("provincia", ""),
        "telefono": request.form["telefono"],
        "correo": request.form.get("correo", ""),
        "solicita": request.form.get("solicita", ""),
        "periodos": request.form.get("periodos", ""),
        "estado": "Pendiente",
    }

    supabase.table("personas").insert(data).execute()
    return redirect("/")


@app.route("/eliminar/<int:persona_id>", methods=["DELETE"])
def eliminar(persona_id):
    persona = obtener_persona_por_id(persona_id)
    if persona:
        archivo = ruta_archivo_generado(persona)
        if archivo.exists():
            archivo.unlink()

    supabase.table("personas").delete().eq("id", persona_id).execute()
    return ("", 204)


@app.route("/generar", methods=["POST"])
def generar():
    payload = request.get_json(silent=True) or {}
    ids = payload.get("ids") or []
    ids = [int(persona_id) for persona_id in ids if str(persona_id).isdigit()]

    if not ids:
        return jsonify({"ok": False, "message": "No hay registros seleccionados."}), 400

    generados = []
    for persona_id in ids:
        persona = obtener_persona_por_id(persona_id)
        if not persona:
            continue
        archivo = generar_y_marcar(persona)
        generados.append({"id": persona_id, "archivo": archivo.name})

    return jsonify({"ok": True, "generados": generados})


@app.route("/descargar/<int:persona_id>")
def descargar(persona_id):
    persona = obtener_persona_por_id(persona_id)
    if not persona:
        return redirect("/")

    archivo = ruta_archivo_generado(persona)
    if not archivo.exists():
        archivo = generar_y_marcar(persona)

    return send_file(archivo, as_attachment=True, download_name=archivo.name)


@app.route("/ver/<int:persona_id>")
def ver(persona_id):
    return redirect(f"/descargar/{persona_id}")


@app.route("/descargar_zip")
def descargar_zip():
    ids_raw = request.args.get("ids", "")
    ids = [int(persona_id) for persona_id in ids_raw.split(",") if persona_id.strip().isdigit()]

    if not ids:
        personas, _, _ = obtener_personas()
        ids = [int(persona["id"]) for persona in personas]

    buffer = BytesIO()
    with ZipFile(buffer, "w", ZIP_DEFLATED) as zip_file:
        for persona_id in ids:
            persona = obtener_persona_por_id(persona_id)
            if not persona:
                continue

            archivo = ruta_archivo_generado(persona)
            if not archivo.exists():
                archivo = generar_y_marcar(persona)

            zip_file.write(archivo, arcname=archivo.name)

    buffer.seek(0)
    nombre_zip = f"oficios_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
    return send_file(buffer, as_attachment=True, download_name=nombre_zip, mimetype="application/zip")


if __name__ == "__main__":
    app.run(debug=True)
