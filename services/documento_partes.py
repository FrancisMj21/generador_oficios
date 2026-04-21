from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from services.documento_layout import (
    ANCHO_COLUMNA_LEMA,
    ANCHO_COLUMNA_LOGO,
    ANCHO_LOGO,
    FUENTE_NOMBRE,
    FUENTE_TAMANIO_LEMA,
    FUENTE_TAMANIO_NORMAL,
    LEMA_1,
    LEMA_2,
    LOGO_PATH,
    PIE_LINEA_1,
    PIE_LINEA_2,
    PIE_LINEA_3,
    PIE_LINEA_4,
)


def agregar_borde_parrafo(parrafo, posicion="bottom"):
    p_pr = parrafo._p.get_or_add_pPr()
    bordes = p_pr.find(qn("w:pBdr"))
    if bordes is None:
        bordes = OxmlElement("w:pBdr")
        p_pr.append(bordes)

    borde = bordes.find(qn(f"w:{posicion}"))
    if borde is None:
        borde = OxmlElement(f"w:{posicion}")
        bordes.append(borde)

    borde.set(qn("w:val"), "single")
    borde.set(qn("w:sz"), "8")
    borde.set(qn("w:space"), "1")
    borde.set(qn("w:color"), "000000")


def agregar_texto(parrafo, texto, bold=False, size=None, underline=False, italic = False):
    run = parrafo.add_run(texto)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.name = FUENTE_NOMBRE
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FUENTE_NOMBRE)
    run.font.size = size or FUENTE_TAMANIO_NORMAL
    return run


def agregar_encabezado_oficio(seccion):
    header = seccion.header
    header.is_linked_to_previous = False

    tabla = header.add_table(rows=1, cols=2, width=ANCHO_COLUMNA_LOGO + ANCHO_COLUMNA_LEMA)
    tabla.alignment = WD_TABLE_ALIGNMENT.CENTER
    tabla.autofit = False
    tabla.columns[0].width = ANCHO_COLUMNA_LOGO
    tabla.columns[1].width = ANCHO_COLUMNA_LEMA

    for celda in tabla.rows[0].cells:
        tc_pr = celda._tc.get_or_add_tcPr()
        tc_borders = tc_pr.first_child_found_in("w:tcBorders")
        if tc_borders is None:
            tc_borders = OxmlElement("w:tcBorders")
            tc_pr.append(tc_borders)
        for borde in ("top", "left", "bottom", "right"):
            elem = OxmlElement(f"w:{borde}")
            elem.set(qn("w:val"), "nil")
            tc_borders.append(elem)

    izquierda_parrafo = tabla.cell(0, 0).paragraphs[0]
    izquierda_parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    izquierda_parrafo.paragraph_format.space_before = Pt(0)
    izquierda_parrafo.paragraph_format.space_after = Pt(0)

    if LOGO_PATH.exists():
        run_logo = izquierda_parrafo.add_run()
        run_logo.add_picture(str(LOGO_PATH), width=ANCHO_LOGO)
    else:
        agregar_texto(izquierda_parrafo, "UGEL TACNA", bold=True)

    derecha = tabla.cell(0, 1).paragraphs[0]
    # derecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    derecha.paragraph_format.space_before = Pt(8)
    derecha.paragraph_format.space_after = Pt(0)
    agregar_texto(derecha, f"{LEMA_1}\n", size=FUENTE_TAMANIO_LEMA)
    agregar_texto(derecha, LEMA_2, size=FUENTE_TAMANIO_LEMA)

    separador = header.add_paragraph()
    separador.paragraph_format.space_before = Pt(1)
    separador.paragraph_format.space_after = Pt(0)
    agregar_borde_parrafo(separador)


def agregar_pie_oficio(seccion):
    footer = seccion.footer
    footer.is_linked_to_previous = False

    seccion.footer_distance = Cm(0.3)
    
    pie = footer.add_paragraph()
    pie.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pie.paragraph_format.space_before = Pt(0)
    pie.paragraph_format.space_after = Pt(0)

    agregar_borde_parrafo(pie, "top")
    agregar_texto(pie, f"{PIE_LINEA_1}\n", size=Pt(6))
    agregar_texto(pie, f"{PIE_LINEA_2}\n", size=Pt(6))
    agregar_texto(pie, f"{PIE_LINEA_3}\n", size=Pt(6))
    agregar_texto(pie, PIE_LINEA_4, size=Pt(6))


def agregar_seccion_info(seccion, persona):

    p = seccion.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1
    agregar_texto(p, "                                            Tacna, ",size=Pt(9))

    # Oficio
    p = seccion.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1
    agregar_texto(p, "OFICIO N°	          -UFESC-URRHH/UGEL.T/GOB.REG.TACNA", bold=True, underline=True, size=Pt(9))


    """Agrega sección de información de persona de forma estructurada."""
    tratamiento = persona.get('tratamiento', '')
    nombre = persona.get('nombre', '').upper()
    cargo = persona.get('cargo', '').upper()
    centro = persona.get('centro_trabajo', '').upper()
    direccion = persona.get('direccion', '')
    provincia = persona.get('provincia', '')
    telefono = persona.get('telefono', '')
    correo = persona.get('correo', '')

    # Tratamiento
    p = seccion.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    agregar_texto(p, f"{tratamiento}: ", bold=True, size=Pt(9))

    # Nombre bold
    p = seccion.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    agregar_texto(p, nombre, bold=True, size=Pt(9))

    # Cargo bold
    p = seccion.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    agregar_texto(p, cargo, size=Pt(9))

    # Centro bold
    p = seccion.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    agregar_texto(p, centro, size=Pt(9))

    # Dirección
    p = seccion.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    agregar_texto(p, f"DIRECCIÓN: {direccion}", size=Pt(9))

    # Provincia
    p = seccion.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    agregar_texto(p, f"Tacna/Tacna/{provincia}", underline=True, bold=True, size=Pt(9))

    # Teléfono
    p = seccion.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    agregar_texto(p, f"Teléfono:", underline=True, bold=True, size=Pt(9))
    agregar_texto(p, f" {telefono}", bold=True, size=Pt(9))

    # Correo
    p = seccion.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    agregar_texto(p, f"Correo electrónico:", underline=True, bold=True , size=Pt(9))
    agregar_texto(p, f" {correo}", bold=True, size=Pt(9))

def agregar_cuerpo_oficio(seccion, data):
    """Agrega cuerpo completo del oficio."""
    solicita = data.get('solicita', '')
    periodos = data.get('periodos', [])
    
    fecha_cud = data.get('fecha_cud', '')
    cud = data.get("numero_cud", "")

    # ASUNTO
    p = seccion.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    agregar_texto(p, "ASUNTO", bold=True, size=Pt(9))
    agregar_texto(p, "\t\t: FORMULO RESPUESTA", bold=True, size=Pt(9))

    # REFERENCIA
    p = seccion.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.space_after = Pt(8)
    agregar_texto(p, "REFERENCIA", bold=True, size=Pt(9))
    agregar_texto(p, f"\t\t: CUD. N° {cud}({fecha_cud})",  size=Pt(9))

    # Separador
    separador = seccion.add_paragraph()
    separador.paragraph_format.space_after = Pt(12)
    agregar_borde_parrafo(separador)

    # Saludo
    saludo = seccion.add_paragraph()
    saludo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    saludo.paragraph_format.first_line_indent = Cm(3)
    saludo.paragraph_format.line_spacing = 1
    saludo.paragraph_format.space_after = Pt(8)
    agregar_texto(
        saludo,
        "Tengo el agrado de dirigirme a usted, para expresarle mi cordial saludo y a la vez brindar "
        f"la debida atención al documento de la referencia, mediante el cual solicita {solicita}.", size=Pt(8)),

    # Cuerpo principal
    cuerpo = seccion.add_paragraph()
    cuerpo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cuerpo.paragraph_format.first_line_indent = Cm(3)
    cuerpo.paragraph_format.line_spacing = 1
    cuerpo.paragraph_format.space_after = Pt(8)
    agregar_texto(
        cuerpo,
        "Al respecto, se deberá tener presente que, el literal l, del artículo 41° de la Ley N° 29944, Ley de la Reforma Magisterial, "
        "expresa que los profesores tienen derecho al reconocimiento de oficio de su tiempo de servicios efectivos, del mismo modo la norma mencionada establece entre sus disposiciones complementarias, transitorias y finales, específicamente en la décima primera que para el cálculo de la asignación por tiempo de servicios se consideran los servicios prestados bajo los regímenes de la Ley N° 24029 y de la Ley 29062, incluyendo el tiempo de servicios prestado en la condición de contratado por servicios personales."
        ,size=Pt(8),
    )

    # Sustento
    sustento = seccion.add_paragraph()
    sustento.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    sustento.paragraph_format.first_line_indent = Cm(3)
    sustento.paragraph_format.line_spacing = 1
    sustento.paragraph_format.space_after = Pt(8)
    agregar_texto(
        sustento,
        "De acuerdo con la R.V.M. Nº 112-2023-MINEDU, de fecha 07 de agosto del 2023 “Disposiciones que regulan los procedimientos técnicos del Escalafón Magisterial” en el sub numeral 6.1.8. se precisa que todo",
        size=Pt(8),
    )
    
    
    agregar_texto(
        sustento,
        "Acto resolutivo que reconoce al profesor nombrado el tiempo de servicio en su condición de contratado en el marco de la LRM, en el inciso a) dice, el profesor puede solicitar el reconocimiento de tiempo de servicio prestado en IIEE públicas, en condición de contratado, ante la DRE/UGEL donde se desempeña actualmente,",
        italic = True,
        size=Pt(8),
    )
    
    agregar_texto(
        sustento,
        " adjuntando las resoluciones de contrato y las boletas o constancias de pago de remuneraciones y descuentos por dichos períodos. ",
        italic = True,
        bold=True,
        size=Pt(8),
    )
    
    agregar_texto(
        sustento,
        "Inciso c) Para el caso de reconocimiento por los servicios prestados en la condición de contratado, se precisa que, no se deben considerar las resoluciones por reconocimiento de pago por periodos menores a treinta (30) días.",
        italic=True,
        size=Pt(8),
    )

    # Cierre
    cierre = seccion.add_paragraph()
    cierre.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cierre.paragraph_format.first_line_indent = Cm(3)
    cierre.paragraph_format.line_spacing = 1
    cierre.paragraph_format.space_after = Pt(6)
    agregar_texto(
        cierre,
        "Al respecto, se procedió a evaluar el expediente de la referencia, en el que se pudo verificar que le falta presentar las boletas y/o constancias de pago de remuneraciones que se indica. Por lo que, se exhorta cumplir con la subsanación de lo indicado para continuar con el procedimiento administrativo. Bajo su responsabilidad.",
        size=Pt(8),
    )

    if periodos:
        periodo_titulo = seccion.add_paragraph()
        periodo_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        periodo_titulo.paragraph_format.space_before = Pt(2)
        periodo_titulo.paragraph_format.space_after = Pt(1)
        agregar_texto(periodo_titulo, "Periodo(s)", bold=True, size=Pt(9))

        for periodo in periodos:
            item = seccion.add_paragraph()
            item.alignment = WD_ALIGN_PARAGRAPH.CENTER
            item.paragraph_format.space_after = Pt(0)
            agregar_texto(item, f"- {periodo}", size=Pt(9))

    # Despedida
    despedida = seccion.add_paragraph()
    despedida.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    despedida.paragraph_format.first_line_indent = Cm(3)
    despedida.paragraph_format.line_spacing = 1
    despedida.paragraph_format.space_before = Pt(5)
    despedida.paragraph_format.space_after = Pt(6)
    agregar_texto(
        despedida,
        "Sin otro particular, hago propicia la oportunidad para reiterarle las muestras de mi especial "
        "consideración y estima personal.",
        size=Pt(8),
    )

    # Firma
    atentamente = seccion.add_paragraph()
    atentamente.alignment = WD_ALIGN_PARAGRAPH.CENTER
    atentamente.paragraph_format.space_after = Pt(10)
    agregar_texto(atentamente, "Atentamente,", size=Pt(9))

    firma = seccion.add_paragraph()
    firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
    firma.paragraph_format.space_after = Pt(28)
    agregar_texto(firma, "GOBIERNO REGIONAL DE TACNA", bold=True, size=Pt(9))

    firma_linea = seccion.add_paragraph()
    firma_linea.alignment = WD_ALIGN_PARAGRAPH.CENTER
    firma_linea.paragraph_format.space_after = Pt(1)
    agregar_texto(firma_linea, "__________________________________________", bold=True, size=Pt(9))

    firma_nombre = seccion.add_paragraph()
    firma_nombre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    firma_nombre.paragraph_format.space_after = Pt(0)
    agregar_texto(firma_nombre, "ABOG. EDICA CARINA CANQUI ATENCIO", bold=True, size=Pt(9))

    firma_cargo = seccion.add_paragraph()
    firma_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    firma_cargo.paragraph_format.space_after = Pt(0)
    agregar_texto(firma_cargo, "JEFE DE LA UNIDAD DE RECURSOS HUMANOS", bold=True, size=Pt(9))

    firma_ugel = seccion.add_paragraph()
    firma_ugel.alignment = WD_ALIGN_PARAGRAPH.CENTER
    firma_ugel.paragraph_format.space_after = Pt(2)
    agregar_texto(firma_ugel, "UGEL TACNA", bold=True, size=Pt(9))

    # Visto
    visto = seccion.add_paragraph()
    visto.paragraph_format.space_before = Pt(6)
    visto.paragraph_format.space_after = Pt(0)
    agregar_texto(visto, "ECCA/UGRH", size=Pt(4))

    archivo = seccion.add_paragraph()
    archivo.paragraph_format.space_after = Pt(2)
    agregar_texto(archivo, "MLBR/AAJPG", size=Pt(4))

